import docx
import pathlib
# Глоссарий общий и спеки, лучше хранить в рахных папках

c = pathlib.Path(r'') #полный путь до папки со спеками
all_files = []
for i in c.iterdir():
    all_files.append(i)



terms = dict()
terms_doc = docx.Document(r'') #полный путь до глоссария
ttt = terms_doc.tables[0]
for row in ttt.rows:
    terms[row.cells[0].text] = row.cells[1].text





for files in all_files:
    doc = docx.Document(files)
    tables = doc.tables
    count = 0
    textt = []
    for paragraph in doc.paragraphs:
        textt.append(paragraph.text)




    # создаем пустой словарь под данные таблиц
    data_tables = {i:None for i in range(len(tables))}
    # проходимся по таблицам
    for i, table in enumerate(tables):
        # создаем список строк для таблицы `i` (пока пустые)
        data_tables[i] = [[] for _ in range(len(table.rows))]
        # проходимся по строкам таблицы `i`
        for j, row in enumerate(table.rows):
            # проходимся по ячейкам таблицы `i` и строки `j`
            for cell in row.cells:
                # добавляем значение ячейки в соответствующий
                # список, созданного словаря под данные таблиц
                data_tables[i][j].append(cell.text)


    q = [i for i in data_tables.values()]
    for i in q:
        for j in i:
            for jj in j:
                textt.append(jj)


    buff = []
    for i in textt:
        if i in terms.keys() and i not in buff:
            if i == 'Комментарий' or i == 'Термин':
                continue
            else:
                buff.append(i)

    for paragraph in doc.paragraphs:
        count += 1
        if count == 8:
            t = doc.add_table(len(buff),2)
            t.style = 'Table Grid'
            for row in range(len(buff)):
                for col in range(2):
                    cell = t.cell(row, col)
                    if col == 0:
                        cell.text = buff[row]
                    else:
                        cell.text = terms[buff[row]]
        elif count > 8:
            break

    doc.save(str(files).split('\\')[-1].split('.')[0]+'_TEST_'+'.docx')


